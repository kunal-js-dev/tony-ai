"""Word document text extraction using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document

logger = logging.getLogger(__name__)


def extract_doc(filepath: str | Path) -> dict[str, Any]:
    """Extract paragraphs and tables from a Word document."""
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    doc = Document(str(path))
    paragraphs: list[str] = []
    tables: list[dict[str, Any]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for t_idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"table_index": t_idx, "rows": rows})
        for row in rows:
            paragraphs.append(" | ".join(row))

    full_text = "\n\n".join(paragraphs)

    return {
        "type": "doc",
        "title": path.stem,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(tables),
        "word_count": len(full_text.split()) if full_text else 0,
        "text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "summary": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(tables),
            "words": len(full_text.split()) if full_text else 0,
        },
    }
